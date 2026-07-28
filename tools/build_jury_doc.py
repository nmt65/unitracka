from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "prezentare-juriu.md"
OUTPUT = ROOT / "docs" / "UniTrack-Ghid-Prezentare-Juriu.docx"

BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
TEAL = "0B9488"
INK = "17313A"
MUTED = "526C72"
PALE = "EAF3F5"
PALE_BLUE = "EAF2F8"
WHITE = "FFFFFF"
CONTENT_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_fixed_table_layout(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def create_numbering_sequence(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"].element.pPr.numPr.numId.val
    source_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_paragraph_border(paragraph, color=BLUE, size=12, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    style_specs = {
        "Title": (28, INK, 0, 12),
        "Subtitle": (13, MUTED, 0, 12),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "UNITRACK  /  INFOEDUCAȚIE 2026"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    set_paragraph_border(header, color="B8D4D8", size=4, space=5)
    add_page_number(section.footer.paragraphs[0])


def add_inline(paragraph, text):
    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BLUE_DARK)
        else:
            paragraph.add_run(token)


def add_cover(doc):
    rule = doc.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    rule.autofit = False
    rule.columns[0].width = Inches(6.5)
    cell = rule.cell(0, 0)
    set_cell_shading(cell, TEAL)
    cell.height = Inches(0.08)
    cell.text = ""

    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_before = Pt(24)
    eyebrow.paragraph_format.space_after = Pt(16)
    run = eyebrow.add_run("PROIECT WEB  ·  INFOEDUCAȚIE 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph(style="Title")
    title.add_run("UniTrack")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Dosar de prezentare pentru juriu")

    statement = doc.add_paragraph()
    statement.paragraph_format.space_before = Pt(16)
    statement.paragraph_format.space_after = Pt(18)
    statement.paragraph_format.line_spacing = 1.18
    run = statement.add_run("Admitere, documente și decizii universitare într-un singur flux verificabil.")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE_DARK)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table_layout(meta)
    widths = [2300, CONTENT_DXA - 2300]
    rows = [
        ("Categorie", "Web"),
        ("Durată", "12 minute + întrebări"),
        ("Aplicație", "https://unitrack.sbs"),
        ("Repository", "https://github.com/nmt65/unitracka"),
    ]
    for row_index, (label, value) in enumerate(rows):
        for col_index, content in enumerate((label, value)):
            cell = meta.cell(row_index, col_index)
            cell.width = Inches(widths[col_index] / 1440)
            set_cell_margins(cell, 90, 120, 90, 120)
            if row_index % 2 == 0:
                set_cell_shading(cell, PALE)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(content)
            run.font.size = Pt(10)
            run.font.bold = col_index == 0
            run.font.color.rgb = RGBColor.from_string(TEAL if col_index == 0 else INK)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    callout.autofit = False
    callout.columns[0].width = Inches(6.5)
    callout_cell = callout.cell(0, 0)
    set_cell_shading(callout_cell, PALE_BLUE)
    set_cell_margins(callout_cell, 180, 220, 180, 220)
    callout_p = callout_cell.paragraphs[0]
    callout_p.paragraph_format.space_before = Pt(4)
    callout_p.paragraph_format.space_after = Pt(4)
    callout_run = callout_p.add_run(
        "Documentul include structura prezentării, discursul complet, arhitectura, "
        "securitatea, funcționalitățile, întrebările juriului și planul de rezervă."
    )
    callout_run.font.size = Pt(11)
    callout_run.font.bold = True
    callout_run.font.color.rgb = RGBColor.from_string(BLUE_DARK)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F6F7")
    set_cell_margins(cell, 130, 160, 130, 160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.3)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_markdown_table(doc, rows):
    if not rows:
        return
    headers = [cell.strip() for cell in rows[0].strip().strip("|").split("|")]
    body_rows = rows[2:] if len(rows) > 1 and set(rows[1].replace("|", "").strip()) <= {"-", ":", " "} else rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_layout(table)
    first_ratio = 0.46 if headers and headers[0].lower() in {"rută", "ruta"} else 0.29
    first_width = int(CONTENT_DXA * first_ratio)
    remaining = CONTENT_DXA - first_width
    widths = [first_width] + [remaining // max(1, len(headers) - 1)] * (len(headers) - 1)

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, value in enumerate(headers):
        cell = header_row.cells[index]
        cell.width = Inches(widths[index] / 1440)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, BLUE_DARK)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

    for row_index, row_text in enumerate(body_rows):
        values = [cell.strip() for cell in row_text.strip().strip("|").split("|")]
        row = table.add_row()
        for index in range(len(headers)):
            value = values[index] if index < len(values) else ""
            cell = row.cells[index]
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2 == 0:
                set_cell_shading(cell, "F7FAFB")
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(8.7)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_markdown(doc, text):
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines = []
    skip_intro = True
    active_numbering_id = None

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if skip_intro and (line.startswith("# UniTrack") or line.startswith("## Dosar de prezentare")):
            index += 1
            continue
        if skip_intro and line.startswith("**Categorie:"):
            while index < len(lines) and lines[index].strip():
                index += 1
            skip_intro = False
            continue
        if line.startswith("## 1."):
            skip_intro = False

        if not line.strip():
            active_numbering_id = None
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].lstrip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            active_numbering_id = None
            level = len(heading.group(1))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            active_numbering_id = None
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            if active_numbering_id is None:
                active_numbering_id = create_numbering_sequence(doc)
            apply_numbering(paragraph, active_numbering_id)
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph_lines = [line.strip()]
        active_numbering_id = None
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith(("#", "-", "```", "|")) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    parse_markdown(doc, markdown)

    properties = doc.core_properties
    properties.title = "UniTrack - Dosar de prezentare pentru juriu"
    properties.subject = "InfoEducație 2026 - secțiunea Web"
    properties.author = "Echipa UniTrack"
    properties.keywords = "UniTrack, InfoEducație, web, admitere, documentație"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
