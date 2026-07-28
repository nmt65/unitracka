from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
BLACK = RGBColor(17, 17, 17)
GRAY = RGBColor(95, 95, 95)
LIGHT = "F1F1EE"


def set_font(run, size=11, bold=None, color=BLACK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_font(run, 9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, before, after in [
        ("Heading 1", 17, 18, 9),
        ("Heading 2", 14, 14, 7),
        ("Heading 3", 12, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "UNITRACK  /  INFOEDUCAȚIE 2026"
    set_font(header.runs[0], 9, True, GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(115)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("UNITRACK")
    set_font(r, 12, True, GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(title)
    set_font(r, 30, True, BLACK, "Aptos Display")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run(subtitle)
    set_font(r, 15, False, GRAY)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(2.15), Inches(2.15), Inches(2.15)]
    values = [("PRODUS", "Platformă web"), ("CATEGORIE", "Web"), ("VERSIUNE", "Națională 2026")]
    for cell, width, (label, value) in zip(table.rows[0].cells, widths, values):
        cell.width = width
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT)
        tc_pr.append(shd)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        a = p.add_run(label + "\n")
        set_font(a, 8, True, GRAY)
        b = p.add_run(value)
        set_font(b, 10.5, True, BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(150)
    r = p.add_run("Admitere universitară organizată, verificabilă și sigură.")
    set_font(r, 12, True, BLACK)
    doc.add_page_break()


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, 9.5, False, BLACK, "Cascadia Mono")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, 11, True)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def add_markdown(doc, path, skip_title=True):
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    seen_title = False
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.right_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_lines))
                set_font(run, 8.5, False, BLACK, "Cascadia Mono")
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            if skip_title and not seen_title:
                seen_title = True
                continue
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, line)


def save_documentation():
    doc = Document()
    configure(doc)
    add_cover(doc, "Documentație tehnică", "Ghid de arhitectură, funcționare, securitate și demonstrație")
    add_markdown(doc, DOCS / "DOCUMENTATIE-PREZENTARE.md")
    doc.core_properties.title = "UniTrack - Documentație pentru juriu"
    doc.core_properties.subject = "InfoEducație 2026"
    doc.save(DOCS / "UniTrack-Documentatie-Juriu.docx")


def save_speeches():
    doc = Document()
    configure(doc)
    add_cover(doc, "Discursuri pentru juriu", "Versiuni cronometrate de 8 și 15 minute")
    doc.add_heading("Versiunea de 8 minute", level=1)
    add_markdown(doc, DOCS / "DISCURS-8-MINUTE.md")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Versiunea de 15 minute", level=1)
    add_markdown(doc, DOCS / "DISCURS-15-MINUTE.md")
    doc.core_properties.title = "UniTrack - Discursuri pentru juriu"
    doc.core_properties.subject = "InfoEducație 2026"
    doc.save(DOCS / "UniTrack-Discursuri-Juriu.docx")


if __name__ == "__main__":
    save_documentation()
    save_speeches()
    print("Documentele au fost generate.")
